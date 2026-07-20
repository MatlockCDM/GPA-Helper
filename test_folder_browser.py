import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from folder_browser import (
    SUPPORTED_EXTENSIONS,
    numbered_folders,
    review_folders,
    searchable_text,
)


class FolderSelectionTests(unittest.TestCase):
    def test_empty_folder_set_has_no_reviews(self):
        self.assertEqual(review_folders({}), [])

    def test_selects_newest_and_available_review_intervals(self):
        folders = {number: f"/{number}" for number in (1, 2, 4, 7)}
        selected = review_folders(folders)
        self.assertEqual([number for number, _ in selected], [7, 4, 2])

    def test_finds_only_folders_with_leading_numbers(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            for name in ("1 first", "20 review", "Lesson 30"):
                (root / name).mkdir()

            found = numbered_folders(root)

        self.assertEqual(sorted(found), [1, 20])

    def test_duplicate_numbers_are_deterministic(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            (root / "1 Alpha").mkdir()
            (root / "1 Beta").mkdir()

            found = numbered_folders(root)

        self.assertEqual(Path(found[1]).name, "1 Alpha")


class DocumentSearchTests(unittest.TestCase):
    def test_supported_extensions_include_common_documents(self):
        self.assertTrue(
            {".txt", ".md", ".csv", ".pdf", ".doc", ".docx", ".odt", ".rtf"}
            <= SUPPORTED_EXTENSIONS
        )

    def test_reads_plain_text_case_independently_of_extension(self):
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "words.MD"
            path.write_text("Salom Dunyo", encoding="utf-8")
            self.assertIn("salom", searchable_text(path).casefold())

    def test_extracts_rtf_text(self):
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "words.rtf"
            path.write_text(r"{\rtf1\ansi Uzbek vocabulary}", encoding="utf-8")
            self.assertIn("Uzbek vocabulary", searchable_text(path))

    def test_extracts_docx_paragraphs_and_tables(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "words.docx"
            document = Document()
            document.add_paragraph("Lesson vocabulary")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "table word"
            document.save(path)
            extracted = searchable_text(path)
            self.assertIn("Lesson vocabulary", extracted)
            self.assertIn("table word", extracted)

    def test_extracts_odt_paragraphs(self):
        from odf.opendocument import OpenDocumentText
        from odf.text import P

        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "words.odt"
            document = OpenDocumentText()
            document.text.addElement(P(text="OpenDocument vocabulary"))
            document.save(path)
            self.assertIn("OpenDocument vocabulary", searchable_text(path))

    def test_extracts_pdf_pages(self):
        fake_module = SimpleNamespace(
            PdfReader=lambda _path: SimpleNamespace(
                pages=[SimpleNamespace(extract_text=lambda: "PDF vocabulary")]
            )
        )
        with patch.dict("sys.modules", {"pypdf": fake_module}):
            self.assertEqual(searchable_text("words.pdf"), "PDF vocabulary")


if __name__ == "__main__":
    unittest.main()
